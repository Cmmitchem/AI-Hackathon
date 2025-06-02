# import os 
# import sys 
# import json 
# import base64
# import requests
# import base64
# import mimetypes
# from pathlib import Path 

# from docx import Document
# import PyPDF2
# import io

# OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

# def extract_text_from_docx(file_path):
#     """Extract text from a word document (.docx)"""
#     try: 
#         doc = Document(file_path)
#         full_text = []

#         #extract text from each paragraph
#         for para in doc.paragraphs: 
#             full_text.append(para.text)

#         for table in doc.tables:
#             for row in table.row: 
#                 for cell in row.cells: 
#                     full_text.append(cell.text)

#         return '\n'.join(full_text)
#     except Exception as e: 
#         raise Exception(f"Error extracting text from word document: {str(e)}")
    
# def extract_text_from_pdf(file_path):
#     """Extract text from a PDF doc"""
#     try: 
#         with open(file_path, 'rb') as file: 
#             pdf_reader = PyPDF2.PdfReader(file)
#             full_text=[]
#             for page_num in range(len(pdf_reader.pages)):
#                 page = pdf_reader.pages[page_num]
#                 full_text.append(page.extract_text())
#             return '\n'.join(full_text)
#     except Exception as e:
#         Exception(f"Error extracting text from PDF: {str(e)}")

# def get_document_text(file_path):
#     """Extract text from document based on file type"""
#     file_extension = Path(file_path).suffix.lower()

#     if file_extension == '.docx': 
#         return extract_text_from_docx(file_path)
#     elif file_extension == '.pdf':
#         return extract_text_from_pdf(file_path)
#     elif file_extension in ['.txt', '.md', '.csv']:
#         # for text files read the content
#         with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
#             return file.read()
        
#     else: 
#         raise Exception(f"unsupported file format: {file_extension}. Supported formats are .docx, .pdf,.txt, .md, and .csv")

# def  get_document_summary(file_path):
#     """Send document to openAI API and get summary response"""
#     if not OPENAI_API_KEY: 
#         raise ValueError("OpenAI API key not found, please set the OPENAI_API_KEY env variable.")
    
#     #get file content 
#     document_text = get_document_text(file_path)
    
#     # prepare the API request
#     headers = {
#         "Authorization": f"Bearer {OPENAI_API_KEY}",
#         "Content-Type": "application/json"
#     }

#     #create message payload 
#     payload = {
#         "model": 'gpt-4o', 
#         "max_tokens": 1000, 
#         "messages": [
#             {
#                 "role": "system",
#                 "content": """I will provide a set of raw notes taken during a meeting, brainstorming session, or research activity. Please analyze the content and organize it into the following four sections:
#                     1.	Summary of the Notes – Provide a concise overview of what the notes cover.
#                     2.	Most Discussed Topics – Identify and list the topics that were mentioned or emphasized most frequently.
#                     3.	Key Points to Know – Extract and clearly present all the important facts, data, or insights conveyed in the notes.
#                     4.	Action Items – List all tasks, follow-ups, or responsibilities that need to be completed, specifying who they are assigned to if possible.
# """
#             },
#             {
#                 "role": "user", 
#                 "content": [
#                     {
#                         "type": "text",
#                         "text": f"""Please rewrite the content using clear, professional language with proper grammar, punctuation, and sentence structure. Do not omit any important details, and preserve the intent and context of the original notes. 
#                         Here are the notes:{document_text}"""
#                     },
                    
#                 ]
#             }
#         ]
#     }

#     # make the API request
#     response = requests.post(
#         "https://api.openai.com/v1/chat/completions", 
#         headers=headers,
#         json=payload
#     )

#     #handle the openAi response
#     if response.status_code != 200:
#         print(f'Error: {response.status_code}')
#         print(response.text)
#         raise Exception(f"API request failed with status code {response.status_code}: {response.text}")
    
#     response_data = response.json()
#     summary = response_data["choices"][0]["message"]["content"]

#     return summary

# def main():
#     """Main function to process the document"""
#     if len(sys.argv) != 2: 
#         print("Usage: python process_document.py <file_path>")
#         sys.exit(1)
    
#     file_path = sys.argv[1]
#     if not os.path.exists(file_path): 
#         print(f"Error: file {file_path} does not exist")
#         sys.exit(1)

#         #check if api key is not set 
#         if not OPENAI_API_KEY: 
#             print(json.dumps({"error": "OpenAI API key not found, please set the OPENAI_API_KEY env variable"}))
#             sys.exit(1)
    
#     try: 
#         #get doc summary from openAI
#         summary = get_document_summary(file_path)

#         print(json.dumps({"summary": summary}))

#     except Exception as e: 
#         print(json.dumps({"error": str(e)}), file=sys.stderr)
#         sys.exit(1)

# if __name__ == "__main__":
#     main()

import os 
import sys 
import json 
import base64
import requests
import base64
import mimetypes
from pathlib import Path 

from docx import Document
import PyPDF2
import io

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

def extract_text_from_docx(file_path):
    """Extract text from a word document (.docx)"""
    try: 
        doc = Document(file_path)
        full_text = []

        #extract text from each paragraph
        for para in doc.paragraphs: 
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows: 
                for cell in row.cells: 
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e: 
        raise Exception(f"Error extracting text from word document: {str(e)}")
    
def extract_text_from_pdf(file_path):
    """Extract text from a PDF doc"""
    try: 
        with open(file_path, 'rb') as file: 
            pdf_reader = PyPDF2.PdfReader(file)
            full_text=[]
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                full_text.append(page.extract_text())
            return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def transcribe_audio_file(file_path):
    """Transcribe audio file using OpenAI Whisper API"""
    try:
        if not OPENAI_API_KEY:
            raise ValueError("OpenAI API key not found")
        
        # Check file size (OpenAI has a 25MB limit for audio files)
        file_size = os.path.getsize(file_path)
        if file_size > 25 * 1024 * 1024:  # 25MB in bytes
            raise Exception("Audio file is too large. Maximum size is 25MB.")
        
        with open(file_path, 'rb') as audio_file:
            # Make request to OpenAI Whisper API
            response = requests.post(
                "https://api.openai.com/v1/audio/transcriptions",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}"
                },
                files={
                    "file": (os.path.basename(file_path), audio_file, "audio/mpeg")
                },
                data={
                    "model": "whisper-1",
                    "response_format": "text"
                }
            )
            
            if response.status_code != 200:
                raise Exception(f"Whisper API request failed with status code {response.status_code}: {response.text}")
            
            transcription = response.text
            return transcription
            
    except Exception as e:
        raise Exception(f"Error transcribing audio file: {str(e)}")

def get_document_text(file_path):
    """Extract text from document based on file type"""
    file_extension = Path(file_path).suffix.lower()

    if file_extension == '.docx': 
        return extract_text_from_docx(file_path)
    elif file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.txt', '.md', '.csv']:
        # for text files read the content
        with open(file_path, 'r', encoding='utf-8', errors="ignore") as file:
            return file.read()
    elif file_extension in ['.mp3', '.wav', '.m4a', '.aac', '.ogg']:
        # for audio files, transcribe them first
        return transcribe_audio_file(file_path)
    else: 
        raise Exception(f"Unsupported file format: {file_extension}. Supported formats are .docx, .pdf, .txt, .md, .csv, .mp3, .wav, .m4a, .aac, .ogg")

def get_document_summary(file_path):
    """Send document to openAI API and get summary response"""
    if not OPENAI_API_KEY: 
        raise ValueError("OpenAI API key not found, please set the OPENAI_API_KEY env variable.")
    
    #get file content 
    document_text = get_document_text(file_path)
    
    # Check if this is an audio file for different system message
    file_extension = Path(file_path).suffix.lower()
    is_audio = file_extension in ['.mp3', '.wav', '.m4a', '.aac', '.ogg']
    
    # prepare the API request
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }

    # Create different system messages for audio vs document content
    if is_audio:
        system_content = """I will provide a transcript of audio content (meeting, brainstorming session, interview, etc.). Please analyze the transcript and organize it into the following four sections:

1. Summary of the Audio Content – Provide a concise overview of what was discussed or covered in the audio.
2. Most Discussed Topics – Identify and list the topics that were mentioned or emphasized most frequently throughout the audio.
3. Key Points to Know – Extract and clearly present all the important facts, insights, decisions, or information conveyed in the audio.
4. Action Items – List all tasks, follow-ups, decisions, or responsibilities that were mentioned, specifying who they are assigned to if possible.

Please clean up any transcription errors and present the information in clear, professional language while preserving all important details and context from the original audio."""
    else:
        system_content = """I will provide a set of raw notes taken during a meeting, brainstorming session, or research activity. Please analyze the content and organize it into the following four sections:

1. Summary of the Notes – Provide a concise overview of what the notes cover.
2. Most Discussed Topics – Identify and list the topics that were mentioned or emphasized most frequently.
3. Key Points to Know – Extract and clearly present all the important facts, data, or insights conveyed in the notes.
4. Action Items – List all tasks, follow-ups, or responsibilities that need to be completed, specifying who they are assigned to if possible."""

    #create message payload 
    payload = {
        "model": 'gpt-4o', 
        "max_tokens": 1000, 
        "messages": [
            {
                "role": "system",
                "content": system_content
            },
            {
                "role": "user", 
                "content": [
                    {
                        "type": "text",
                        "text": f"""Please rewrite the content using clear, professional language with proper grammar, punctuation, and sentence structure. Do not omit any important details, and preserve the intent and context of the original {'transcript' if is_audio else 'notes'}. 
                        
Here {'is the transcript' if is_audio else 'are the notes'}: {document_text}"""
                    },
                    
                ]
            }
        ]
    }

    # make the API request
    response = requests.post(
        "https://api.openai.com/v1/chat/completions", 
        headers=headers,
        json=payload
    )

    #handle the openAi response
    if response.status_code != 200:
        print(f'Error: {response.status_code}')
        print(response.text)
        raise Exception(f"API request failed with status code {response.status_code}: {response.text}")
    
    response_data = response.json()
    summary = response_data["choices"][0]["message"]["content"]

    return summary

def main():
    """Main function to process the document"""
    if len(sys.argv) != 2: 
        print("Usage: python process_document.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path): 
        print(f"Error: file {file_path} does not exist")
        sys.exit(1)

    #check if api key is not set 
    if not OPENAI_API_KEY: 
        print(json.dumps({"error": "OpenAI API key not found, please set the OPENAI_API_KEY env variable"}))
        sys.exit(1)
    
    try: 
        #get doc summary from openAI
        summary = get_document_summary(file_path)

        print(json.dumps({"summary": summary}))

    except Exception as e: 
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()